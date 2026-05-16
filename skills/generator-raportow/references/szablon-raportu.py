"""
szablon-raportu.py — SZABLON generatora raportu akademickiego.

SKOPIUJ ten plik jako bazę nowego skryptu raportu, NIE edytuj go w miejscu.
Cel docelowy: <lokalizacja wskazana przez użytkownika>/_build_raport_XX.py

Pipeline: dane liczbowe w kodzie -> obliczenia w Pythonie -> raport .docx
(opcjonalnie + zalacznik .xlsx lub osadzone wykresy .png). Raport regenerowalny
jednym poleceniem: python <skrypt>.py

Uruchomiony bez zmian generuje szkieletowy raport w katalogu tymczasowym —
to test, ze helpery dzialaja w srodowisku.

Wymaga: python-docx. Opcjonalnie: matplotlib (wykresy), openpyxl (zalacznik xlsx),
numpy (obliczenia). Instalacja: pip install --user python-docx matplotlib openpyxl numpy
"""

import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# matplotlib i openpyxl — usun te importy, jesli raport ich nie uzywa
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side


# ===========================================================================
# KONFIGURACJA STYLU — ustaw raz pod konwencje przedmiotu
# ===========================================================================
# Wariant domyslny: Times New Roman 12, interlinia 1.5, wciecie 0.75 cm
# Wariant alternatywny: FONT="Calibri", ROZMIAR=11, INTERLINIA=1.0,
#                       WCIECIE=0.0, KOLOR_NAGLOWKA=(0x1F, 0x4E, 0x78)
FONT = "Times New Roman"
ROZMIAR = 12
INTERLINIA = 1.5
WCIECIE = 0.75                  # wciecie pierwszego wiersza akapitu [cm]
KOLOR_NAGLOWKA = None           # None = czarny; albo krotka RGB, np. (0x1F, 0x4E, 0x78)
KOLOR_TABELI = "4F81BD"         # tlo wiersza naglowka tabeli (hex)


# ===========================================================================
# DOKUMENT + HELPERY DOCX
# ===========================================================================
doc = Document()


def setup_dokument():
    """Marginesy 2.5 cm i styl Normal wg KONFIGURACJI."""
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(2.5)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(ROZMIAR)
    st.paragraph_format.line_spacing = INTERLINIA
    st.paragraph_format.first_line_indent = Cm(WCIECIE)
    st.paragraph_format.space_after = Pt(0)


def add_heading(text, level=1, size=14, page_break=False):
    """Naglowek sekcji. Rozmiar maleje o 2 pt na kazdy poziom zaglebienia."""
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size - 2 * (level - 1))
    r.bold = True
    if KOLOR_NAGLOWKA:
        r.font.color.rgb = RGBColor(*KOLOR_NAGLOWKA)
    return p


def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(ROZMIAR)
    r.bold = bold
    r.italic = italic
    return p


def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.italic = True
    r.font.size = Pt(ROZMIAR)
    return p


def _cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, col_widths=None, font_size=10):
    """Tabela z tlem wiersza naglowka. rows = lista list (komorki jako str)."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(str(h))
        r.font.name = FONT
        r.font.size = Pt(font_size)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _cell_bg(c, KOLOR_TABELI)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(str(val))
            r.font.name = FONT
            r.font.size = Pt(font_size)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for j, w in enumerate(col_widths):
            for c in t.columns[j].cells:
                c.width = w
    return t


def caption(text):
    """Podpis nad tabela/wykresem (np. 'Tabela 1. ...')."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(ROZMIAR - 1)
    r.bold = True
    return p


def source(text="Źródło: opracowanie własne."):
    """Adnotacja zrodla pod tabela/wykresem."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(ROZMIAR - 2)
    r.italic = True
    return p


def add_image(path, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))


def strona_tytulowa(tytul, podtytul="", naglowek="", prowadzacy="", zespol=None, miejsce_rok=""):
    """Strona tytulowa + podzial strony. zespol = lista nazwisk albo None."""
    for _ in range(4):
        doc.add_paragraph()
    if naglowek:
        add_para(naglowek, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    t = add_para(tytul, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    t.runs[0].font.size = Pt(20)
    if podtytul:
        s = add_para(podtytul, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        s.runs[0].font.size = Pt(13)
    for _ in range(6):
        doc.add_paragraph()
    if prowadzacy:
        add_para(prowadzacy, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    if zespol:
        doc.add_paragraph()
        add_para(" · ".join(zespol), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    if miejsce_rok:
        doc.add_paragraph()
        add_para(miejsce_rok, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_page_break()


def spis_tresci(pozycje):
    """Statyczny spis tresci. pozycje = lista stringow; podsekcje wcinaj spacjami."""
    add_heading("Spis treści", level=1)
    for poz in pozycje:
        p = doc.add_paragraph(poz)
        p.paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()


def wykaz(naglowek, pozycje, numerowany=True):
    """Bibliografia / wykaz tabel / wykaz wykresow jako lista."""
    add_heading(naglowek, level=1)
    styl = "List Number" if numerowany else "List Bullet"
    for poz in pozycje:
        doc.add_paragraph(str(poz), style=styl)


def fmt(x, nd=2):
    """Liczba w zapisie PL — przecinek dziesietny."""
    if isinstance(x, float):
        return f"{x:.{nd}f}".replace(".", ",")
    return str(x)


# ===========================================================================
# HELPERY MATPLOTLIB  (jesli raport osadza wykresy PNG)
# ===========================================================================
def setup_matplotlib():
    matplotlib.rcParams["font.family"] = "DejaVu Sans"
    matplotlib.rcParams["axes.unicode_minus"] = False


def zapisz_wykres(nazwa, folder):
    """Zapisz biezaca figure matplotlib do folder/nazwa (170 dpi). Zwraca sciezke."""
    folder = Path(folder)
    folder.mkdir(exist_ok=True)
    sciezka = folder / nazwa
    plt.savefig(sciezka, dpi=170, bbox_inches="tight")
    plt.close()
    return sciezka


# ===========================================================================
# HELPERY OPENPYXL  (jesli raport ma zalacznik .xlsx z obliczeniami)
# ===========================================================================
_THIN = Side(style="thin", color="999999")
XLSX_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
XLSX_HEAD_FILL = PatternFill("solid", fgColor=KOLOR_TABELI)


def xlsx_header(cell):
    cell.fill = XLSX_HEAD_FILL
    cell.font = Font(bold=True, color="FFFFFF", size=10)
    cell.alignment = Alignment(horizontal="center", vertical="center")
    cell.border = XLSX_BORDER


def xlsx_data(cell, number_format=None, bold=False):
    cell.border = XLSX_BORDER
    cell.alignment = Alignment(horizontal="center", vertical="center")
    if number_format:
        cell.number_format = number_format
    if bold:
        cell.font = Font(bold=True)


# ===========================================================================
# RAPORT — WYPELNIJ
# ===========================================================================
# Ponizej minimalny szkielet. Po skopiowaniu szablonu zastap sekcje DANE
# i SEKCJE MERYTORYCZNE realna zawartoscia, w kolejnosci z checklista-raportu.md.

if __name__ == "__main__":
    setup_dokument()

    # --- DANE WEJSCIOWE — WYPELNIJ ---------------------------------------
    # Tu wstaw tablice danych i funkcje obliczeniowe specyficzne dla raportu.

    # --- STRONA TYTULOWA -------------------------------------------------
    strona_tytulowa(
        tytul="Tytuł raportu",
        podtytul="Podtytuł / firma / system pracy",
        naglowek="NAZWA PRZEDMIOTU — RODZAJ ZAJĘĆ",
        prowadzacy="Prowadzący: <imię i nazwisko>",
        zespol=["Autor 1", "Autor 2"],
        miejsce_rok="Miasto, 2026",
    )

    # --- SPIS TRESCI -----------------------------------------------------
    spis_tresci(["1. Wstęp", "2. <sekcja merytoryczna>", "Bibliografia"])

    # --- SEKCJE MERYTORYCZNE — WYPELNIJ ----------------------------------
    add_heading("1. Wstęp", level=1)
    add_para("Treść wstępu raportu — cel, kontekst, powiązanie z poprzednimi raportami.")
    add_formula("F = a + b · t")
    caption("Tabela 1. Przykładowa tabela.")
    add_table(["Kolumna A", "Kolumna B"], [["wiersz", fmt(12.5)]])
    source()

    setup_matplotlib()
    tmp = Path(tempfile.gettempdir())
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(["A", "B", "C"], [3, 5, 2])
    ax.set_title("Wykres przykladowy")
    png = zapisz_wykres("_szablon-wykres.png", tmp)
    caption("Wykres 1. Przykładowy wykres.")
    add_image(png, width_cm=12)
    source()

    wykaz("Bibliografia", ["Nazwisko, I. (rok). Tytuł. Wydawnictwo."])

    out = tmp / "_szablon-test.docx"
    doc.save(out)
    print(f"OK DOCX -> {out}")

    # --- ZALACZNIK XLSX (opcjonalny) — przyklad uzycia helperow openpyxl --
    wb = Workbook()
    ws = wb.active
    ws.title = "Dane"
    for j, h in enumerate(["t", "A[t]"], start=1):
        xlsx_header(ws.cell(1, j, h))
    for i, val in enumerate([100, 120, 95], start=2):
        xlsx_data(ws.cell(i, 1, i - 1), number_format="0")
        xlsx_data(ws.cell(i, 2, val), number_format="0")
    xlsx_out = tmp / "_szablon-test.xlsx"
    wb.save(xlsx_out)
    print(f"OK XLSX -> {xlsx_out}")
